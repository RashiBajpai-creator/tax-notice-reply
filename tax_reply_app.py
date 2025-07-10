
import streamlit as st
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF
import io

def extract_text_from_pdf(file):
    reader = PdfReader(file)
    return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_summary_from_excel(file):
    df = pd.read_excel(file, engine='openpyxl')
    return df.head(5).to_string(index=False)

def generate_word_reply(notice_text, data_analysis, prev_format_text):
    doc = Document()
    doc.add_paragraph("Reply to Notice\n")
    doc.add_paragraph("Subject: Response to Income Tax Notice\n")

    doc.add_paragraph("Based on the analysis of uploaded documents, here is the response:")
    doc.add_paragraph("\n--- Notice Reference ---\n" + notice_text[:1000])
    doc.add_paragraph("\n--- Summary from PDFs/Excels ---\n" + data_analysis[:2000])
    doc.add_paragraph("\n--- Reply Format (Based on Previous Submission) ---\n" + prev_format_text[:2000])
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf_reply(content):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    for line in content.split('\n'):
        pdf.multi_cell(0, 10, line)
    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer

st.title("Auto Reply to Tax Notice")

st.header("Step 1: Upload PDFs and Excel Files for Analysis")
pdf_excel_files = st.file_uploader("Upload PDF/Excel Files", type=["pdf", "xls", "xlsx"], accept_multiple_files=True)

st.header("Step 2: Upload the Notice (PDF or DOCX)")
notice_file = st.file_uploader("Upload Notice File", type=["pdf", "docx"])

st.header("Step 3: Upload Previous Submission Format (DOCX)")
prev_format_file = st.file_uploader("Upload Previous Submission", type=["docx"])

if st.button("Generate Reply"):
    st.subheader("Generating reply...")
    data_analysis = ""
    for f in pdf_excel_files:
        if f.name.endswith(".pdf"):
            data_analysis += f"\n--- {f.name} ---\n" + extract_text_from_pdf(f)
        elif f.name.endswith((".xls", ".xlsx")):
            data_analysis += f"\n--- {f.name} ---\n" + extract_summary_from_excel(f)

    if notice_file.name.endswith(".pdf"):
        notice_text = extract_text_from_pdf(notice_file)
    else:
        notice_text = extract_text_from_docx(notice_file)

    prev_format_text = extract_text_from_docx(prev_format_file) if prev_format_file else ""

    word_doc = generate_word_reply(notice_text, data_analysis, prev_format_text)
    st.download_button("Download Reply as Word", word_doc, file_name="Reply.docx")

    pdf_content = f"Notice Summary:\n{notice_text}\n\nData Summary:\n{data_analysis}\n\nReply Based on Format:\n{prev_format_text}"
    pdf_doc = generate_pdf_reply(pdf_content)
    st.download_button("Download Reply as PDF", pdf_doc, file_name="Reply.pdf")

    st.success("Reply successfully generated.")
